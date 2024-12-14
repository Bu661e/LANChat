from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # 标题
    title = doc.add_heading('基于Python的局域网聊天室实验报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 实验目的
    doc.add_heading('1. 实验目的', level=1)
    objectives = [
        '掌握基于Socket的网络编程技术',
        '理解TCP/IP协议在实际应用中的应用',
        '实现一个具有图形用户界面的局域网聊天室应用'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    # 2. 实验环境
    doc.add_heading('2. 实验环境', level=1)
    doc.add_paragraph('操作系统：Windows')
    doc.add_paragraph('开发语言：Python')
    doc.add_paragraph('主要库：')
    libraries = [
        'socket：网络通信',
        'PySide6：图形用户界面',
        'threading：多线程处理',
        'json：数据序列化',
        'datetime：时间处理'
    ]
    for lib in libraries:
        p = doc.add_paragraph(lib, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 3. 系统架构
    doc.add_heading('3. 系统架构', level=1)
    
    # 3.1 服务器端
    doc.add_heading('3.1 服务器端（Server）', level=2)
    server_features = [
        '采用TCP Socket实现',
        '支持多客户端并发连接',
        '使用多线程处理客户端消息',
        '实现了用户管理、消息转发等功能'
    ]
    for feature in server_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 3.2 客户端
    doc.add_heading('3.2 客户端（Client）', level=2)
    client_features = [
        '图形用户界面（GUI）实现',
        '支持用户登录',
        '提供群聊（广场）和私聊功能',
        '支持多媒体消息（图片、视频、音频、文件）'
    ]
    for feature in client_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 4. 主要功能
    doc.add_heading('4. 主要功能', level=1)
    # 4.1 用户管理
    doc.add_heading('4.1 用户管理', level=2)
    user_management = [
        '用户登录/登出',
        '在线用户列表维护',
        '用户状态广播'
    ]
    for feature in user_management:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 4.2 消息通信
    doc.add_heading('4.2 消息通信', level=2)
    message_features = [
        '群聊消息',
        '私聊消息',
        '图片消息',
        '视频消息',
        '音频消息',
        '文件传输'
    ]
    for feature in message_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 5. 关键实现
    doc.add_heading('5. 关键实现', level=1)
    
    # 5.1 服务器端实现
    doc.add_heading('5.1 服务器端实现', level=2)
    doc.add_paragraph('1. 服务器初始化')
    p = doc.add_paragraph()
    code = '''def start(self, host='192.168.31.227', port=8000):
    self.server_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    self.server_socket.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
    self.server_socket.bind((self.host, self.port))
    self.server_socket.listen(5)'''
    p.add_run(code).font.name = 'Courier New'
    
    doc.add_paragraph('2. 客户端连接处理')
    p = doc.add_paragraph()
    code = '''def accept_connections(self):
    while self.running:
        client_socket, address = self.server_socket.accept()
        client_thread = threading.Thread(
            target=self.receive_messages,
            args=(client_socket, address)
        )
        client_thread.daemon = True
        client_thread.start()'''
    p.add_run(code).font.name = 'Courier New'
    
    # 5.2 客户端实现
    doc.add_heading('5.2 客户端实现', level=2)
    doc.add_paragraph('1. 连接服务器')
    p = doc.add_paragraph()
    code = '''def connect_to_server(self, server_ip, server_port, local_ip, local_port, username):
    self.socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    self.socket.bind((self.local_ip, self.local_port))
    self.socket.connect((self.server_ip, self.server_port))'''
    p.add_run(code).font.name = 'Courier New'
    
    doc.add_paragraph('2. 消息发送')
    p = doc.add_paragraph()
    code = '''def send_message(self, message):
    message_json = json.dumps(message)
    message_length = len(message_json.encode())
    length_bytes = message_length.to_bytes(4, 'big')
    self.socket.sendall(length_bytes)
    self.socket.sendall(message_json.encode())'''
    p.add_run(code).font.name = 'Courier New'
    
    # 6. 通信协议设计
    doc.add_heading('6. 通信协议设计', level=1)
    
    # 6.1 消息格式
    doc.add_heading('6.1 消息格式', level=2)
    doc.add_paragraph('所有消息采用JSON格式，包含以下字段：')
    message_fields = [
        'type: 消息类型',
        'username: 发送者用户名',
        'content: 消息内容',
        'timestamp: 时间戳',
        '其他特定字段（根据消息类型）'
    ]
    for field in message_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    # 6.2 消息类型
    doc.add_heading('6.2 消息类型', level=2)
    doc.add_paragraph('1. 系统消息')
    system_messages = [
        'login: 用户登录',
        'logout: 用户登出',
        'new_friend_login: 新用户上线通知',
        'one_user_logout: 用户下线通知'
    ]
    for msg in system_messages:
        p = doc.add_paragraph(msg, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph('2. 聊天消息')
    chat_messages = [
        'square_message: 群聊文本消息',
        'private_message: 私聊文本消息',
        'square_image: 群聊图片消息',
        'private_image: 私聊图片消息',
        'square_video: 群聊视频消息',
        'private_video: 私聊视频消息',
        'square_audio: 群聊音频消息',
        'private_audio: 私聊音频消息',
        'square_file: 群聊文件消息',
        'private_file: 私聊文件消息'
    ]
    for msg in chat_messages:
        p = doc.add_paragraph(msg, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 7. 实验总结
    doc.add_heading('7. 实验总结', level=1)
    
    # 7.1 实验成果
    doc.add_heading('7.1 实验成果', level=2)
    doc.add_paragraph('成功实现了一个功能完整的局域网聊天室系统，包括：')
    achievements = [
        '稳定的服务器-客户端通信架构',
        '直观的图形用户界面',
        '丰富的消息类型支持',
        '可靠的文件传输功能'
    ]
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    # 7.2 技术要点
    doc.add_heading('7.2 技术要点', level=2)
    technical_points = [
        '使用TCP Socket确保可靠的数据传输',
        '采用多线程处理提高并发性能',
        '实现了完整的消息协议',
        '支持大文件的分块传输'
    ]
    for point in technical_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # 7.3 可改进之处
    doc.add_heading('7.3 可改进之处', level=2)
    improvements = [
        '添加数据加密功能',
        '实现消息持久化存储',
        '优化大文件传输性能',
        '添加用户认证机制',
        '实现服务器集群部署'
    ]
    for improvement in improvements:
        doc.add_paragraph(improvement, style='List Bullet')
    
    # 保存文档
    doc.save('report.docx')

if __name__ == '__main__':
    create_report()
